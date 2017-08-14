# -*- coding: utf-8 -*-
'''
Created on 2017-07-11 14:39
---------
@summary:
---------
@author: Boris
'''

from docx import Document
from db.oracledb import OracleDB

oracle = OracleDB()


# dir_docx = 'test.docx'
# document = Document(dir_docx)
# for p in document.paragraphs:
#     print(p.text)


def parser_docx_paragraphs(docx):
    document = Document(dir_docx)
    for p in document.paragraphs:
        print(p.text)

def parse_docx_table(f):
    person = {}

    document = Document(f)
    table = document.tables[0]

    for i in range(109):
        name = table.cell(i, 1).text.strip()
        origation = table.cell(i, 5).text.strip()

        if origation not in person.keys():
            person[origation] = set()

        person[origation].add(name)

        # print (name, origation)

    from pprint import pprint
    pprint(person)

    filename = "persons.txt"
    # with open(filename, mode = 'w', encoding = 'utf-8') as file:
    #     for origation, names in person.items():
    #         print(origation)
    #         print(name)
    #         content = origation + ':\n' + ','.join(names)
    #         file.write(content)
    #         file.write('\n------------------------------------\n')

    sequence = 1
    for origation, names in person.items():
        print(origation)

        # sql = 'insert into TAB_IOPM_FIRST_CLUES_CLASSIFY t (t.first_classify_id, t.first_classify, t.zero_id) values (%s, \'%s\', 1)'%(sequence, '')
        # oracle.add(sql)

        for name in names:
            print(name)
            sql = 'insert into TAB_IOPM_CLUES (id, name, Keyword2, First_Id) values (%s, \'%s\', \'%s\', %s)'%("sequence.nextval", name, name, 1)
            oracle.add(sql)
        sequence+=1

# parse_docx_table('test.docx')

jigous = '中华人民共和国国家新闻出版广电总局,中央人民广播电台,中央电视台,中国国际广播电台,国家新闻出版广电总局广播电视卫星直播管理中心,国家新闻出版广电总局广播电视规划院,中广电广播电影电视设计研究院,国家新闻出版广电总局电影剧本规划策划中心,国家新闻出版广电总局电影数字节目管理中心,国家新闻出版广电总局机关服务,国家新闻出版广电总局研修学院（培训中心）,中国广播电影,视社会组织联合会,中央新影集团,中国新闻出版研究院,新闻出版广电总局培训中心,新闻出版总署出版产品质量监督检测中心,中国版本图书馆（国家新闻出版广电总局出版物数据中心）,国家出版基金规划管理办公室,中国印刷博物馆,天津市文化广播影视局,河北省新闻出版广电局,山西省新闻出版广电局,新疆维吾尔自治区广播电影电视局,宁夏回族自治区新闻出版广电局,青海省广播电影电视局,甘肃省新闻出版广电局,陕西省新闻出版广电局,西藏自治区新闻出版广电局,云南省新闻出版局,贵州省新闻出版广电局,四川省新闻出版广电局,重庆市文化委员会,海南省文化广电出版体育厅,广西壮族自治区新闻出版广电局,广东省新闻出版广电局,湖南省新闻出版广电局,湖北省新闻出版广电局湖北省版权局,江西省新闻出版广电局（江西省版权局）,山东省新闻出版广电局,辽宁省新闻出版广电局,吉林省省新闻出版广电局,黑龙江省新闻出版广电局 黑龙江省版权局,福建省新闻出版广电（版权）局,上海市文化广播影视管理局,江苏省新闻出版广电局（版权局）,浙江省新闻出版广电局（省版权局）,安徽省新闻出版广电局'.split(',')

for name in jigous:
    sql = 'insert into TAB_IOPM_CLUES (id, name, Keyword2, First_Id, zero_id) values (%s, \'%s\', \'%s\', %s, %s)'%("sequence.nextval", name, name, 6, 3)
    oracle.add(sql)